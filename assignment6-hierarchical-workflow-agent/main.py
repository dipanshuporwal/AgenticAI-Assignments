import operator
from typing import TypedDict, Sequence, Annotated
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import PromptTemplate
from langgraph.graph import StateGraph, END
from docx import Document
from dotenv import load_dotenv
from IPython.display import Image, display
import os

# Load environment
load_dotenv()


# === State Definition ===
class AgentState(TypedDict):
    messages: Annotated[Sequence[str], operator.add]
    topic: str  # e.g., "medical", "financial"
    summary: str


# === LLM ===
llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash")


# === Supervisor Node ===
def supervisor_node(state: AgentState) -> AgentState:
    query = state["messages"][-1].lower()
    print(f"[Supervisor] Received: {query}")

    if "health" in query or "disease" in query or "treatment" in query:
        topic = "medical"
    elif "stock" in query or "finance" in query or "market" in query:
        topic = "financial"
    else:
        topic = "general"

    return {"messages": [f"Topic classified as: {topic}"], "topic": topic}


# === Research Nodes ===
def medical_research_node(state: AgentState) -> AgentState:
    query = state["messages"][0]
    prompt = PromptTemplate.from_template(
        "Research this medical topic: {query}"
    )
    response = (prompt | llm | StrOutputParser()).invoke({"query": query})
    return {"messages": [f"[Medical Research] {response}"]}


def financial_research_node(state: AgentState) -> AgentState:
    query = state["messages"][0]
    prompt = PromptTemplate.from_template(
        "Research this financial topic: {query}"
    )
    response = (prompt | llm | StrOutputParser()).invoke({"query": query})
    return {"messages": [f"[Financial Research] {response}"]}


# === Report Nodes ===
def create_summary_node(state: AgentState) -> AgentState:
    research_text = state["messages"][-1]
    prompt = PromptTemplate.from_template(
        "Create a short summary from: {research_text}"
    )
    summary = (prompt | llm | StrOutputParser()).invoke(
        {"research_text": research_text}
    )
    return {"messages": [f"[Summary] {summary}"], "summary": summary}


def save_summary_node(state: AgentState) -> AgentState:
    summary = state["summary"]
    document = Document()
    document.add_heading("Research Summary", 0)
    document.add_paragraph(summary)
    os.makedirs("output", exist_ok=True)
    filepath = f"output/summary_{summary[0:5]}.docx"
    document.save(filepath)
    return {"messages": [f"[Document Saved] Summary saved to {filepath}"]}


# === Routers ===
def research_router(state: AgentState) -> str:
    return (
        "MedicalResearch"
        if state["topic"] == "medical"
        else "FinancialResearch"
    )


# === Workflow ===
graph = StateGraph(AgentState)

# Nodes
graph.add_node("Supervisor", supervisor_node)
graph.add_node("MedicalResearch", medical_research_node)
graph.add_node("FinancialResearch", financial_research_node)
graph.add_node("CreateSummary", create_summary_node)
graph.add_node("SaveSummary", save_summary_node)

# Entry
graph.set_entry_point("Supervisor")

# Supervisor → Research
graph.add_conditional_edges(
    "Supervisor",
    research_router,
    {
        "MedicalResearch": "MedicalResearch",
        "FinancialResearch": "FinancialResearch",
    },
)

# Research → Report
graph.add_edge("MedicalResearch", "CreateSummary")
graph.add_edge("FinancialResearch", "CreateSummary")

# Report Steps
graph.add_edge("CreateSummary", "SaveSummary")
graph.add_edge("SaveSummary", END)

# Compile
app = graph.compile()

# ========== Display App Workflow==========
display(Image(app.get_graph().draw_mermaid_png()))

# === Example Run ===
if __name__ == "__main__":
    query = "What are the latest treatments for type 2 diabetes?"
    state = {"messages": [query], "topic": "", "summary": ""}
    result = app.invoke(state)
    print("\n📄 Final Output:")
    for msg in result["messages"]:
        print(msg)
